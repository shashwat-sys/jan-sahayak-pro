from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/apple/Documents/New project/output/documents")
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "0B2545"
BLUE = "1F4D78"
TEAL = "0F766E"
GOLD = "B8892E"
GOLD_DARK = "7A5A00"
SLATE = "334155"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "F8F1E3"
LIGHT_GRAY = "F2F4F7"
BORDER = "CBD5E1"
WHITE = "FFFFFF"
BLACK = "111827"
RED = "9B1C1C"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run(run, font="Aptos", size=None, color=None, bold=None, italic=None, caps=False):
    run.font.name = font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if caps:
        rpr = run._element.get_or_add_rPr()
        if rpr.find(qn("w:caps")) is None:
            rpr.append(OxmlElement("w:caps"))


def set_paragraph(p, before=0, after=6, line=1.1, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def clear_paragraph(p):
    for run in list(p.runs):
        run._element.getparent().remove(run._element)


def paragraph_border_bottom(p, color=GOLD, size="12", space="4"):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_table_layout(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(sum(widths)))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(widths[i]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_cell(cell, font="Aptos", size=9.5, color=BLACK, bold=False, align=None):
    for p in cell.paragraphs:
        set_paragraph(p, before=0, after=0, line=1.08, align=align)
        for run in p.runs:
            set_run(run, font=font, size=size, color=color, bold=bold)


def set_core_properties(doc, title):
    cp = doc.core_properties
    cp.title = title
    cp.author = "16Arena"
    cp.last_modified_by = "16Arena"
    cp.subject = "16Arena compliance documentation"
    cp.keywords = "16Arena, PROG Act, OGAI, compliance"


def setup_doc(title, premium=False):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78 if premium else 0.88)
    section.bottom_margin = Inches(0.78 if premium else 0.88)
    section.left_margin = Inches(0.82 if premium else 0.9)
    section.right_margin = Inches(0.82 if premium else 0.9)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    set_core_properties(doc, title)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 15, NAVY, 14, 6),
        ("Heading 2", 12.5, BLUE, 10, 4),
        ("Heading 3", 11, SLATE, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    return doc


def running_header_footer(doc, left, right, confidentiality="Privileged and Confidential"):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    clear_paragraph(p)
    set_paragraph(p, after=2, line=1.0)
    r = p.add_run(left)
    set_run(r, size=8.5, color=MUTED, bold=True, caps=True)
    r = p.add_run("    |    " + confidentiality)
    set_run(r, size=8.5, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    p = footer.paragraphs[0]
    clear_paragraph(p)
    set_paragraph(p, after=0, line=1.0)
    r = p.add_run(right)
    set_run(r, size=8, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_number(p):
    r = p.add_run("Page ")
    set_run(r, size=8, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_end)


def add_masthead(doc, label, title, subtitle, meta_rows, accent=GOLD):
    p = doc.add_paragraph()
    set_paragraph(p, before=2, after=3, line=1.0)
    r = p.add_run(label.upper())
    set_run(r, size=9, color=accent, bold=True, caps=True)

    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=2, line=1.0)
    r = p.add_run(title)
    set_run(r, font="Aptos Display", size=22, color=NAVY, bold=True)

    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=12, line=1.08)
    r = p.add_run(subtitle)
    set_run(r, size=11.5, color=SLATE)

    table = doc.add_table(rows=len(meta_rows), cols=2)
    set_table_layout(table, [1700, 7660])
    set_table_borders(table, color="D8DEE8", size="4")
    for i, (k, v) in enumerate(meta_rows):
        c0, c1 = table.rows[i].cells
        set_cell_shading(c0, LIGHT_BLUE if i % 2 == 0 else LIGHT_GRAY)
        set_cell_shading(c1, WHITE)
        c0.text = k
        c1.text = v
        format_cell(c0, size=8.8, color=NAVY, bold=True)
        format_cell(c1, size=9.2, color=BLACK)

    p = doc.add_paragraph()
    set_paragraph(p, after=8)
    paragraph_border_bottom(p, color=accent, size="10", space="2")


def add_letterhead(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_layout(table, [180, 6620, 2560])
    set_table_borders(table, color=WHITE, size="0")
    left, main, side = table.rows[0].cells
    set_cell_shading(left, NAVY)
    set_cell_shading(main, WHITE)
    set_cell_shading(side, LIGHT_GOLD)
    main.text = ""
    side.text = ""

    p = main.paragraphs[0]
    set_paragraph(p, after=0, line=1.0)
    r = p.add_run("SHASHWAT")
    set_run(r, font="Aptos Display", size=24, color=NAVY, bold=True, caps=True)
    p = main.add_paragraph()
    set_paragraph(p, after=1, line=1.0)
    r = p.add_run("Advocate | Patna High Court and Supreme Court of India")
    set_run(r, size=8.8, color=SLATE, bold=True)
    p = main.add_paragraph()
    set_paragraph(p, after=0, line=1.0)
    r = p.add_run("LL.B, Faculty of Law, University of Delhi | B.A. (Hons.) History, Kirori Mal College")
    set_run(r, size=8, color=MUTED)

    p = side.paragraphs[0]
    set_paragraph(p, after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("COUNSEL'S REPRESENTATION")
    set_run(r, size=8, color=GOLD_DARK, bold=True, caps=True)
    p = side.add_paragraph()
    set_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("Privileged | Confidential")
    set_run(r, size=8.5, color=NAVY, bold=True)

    p = doc.add_paragraph()
    set_paragraph(p, after=10)
    paragraph_border_bottom(p, color=GOLD, size="12", space="2")


def add_callout(doc, label, text, fill=LIGHT_BLUE, border=GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_layout(table, [9360])
    set_table_borders(table, color=border, size="6")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, after=3, line=1.08)
    r = p.add_run(label.upper())
    set_run(r, size=8.5, color=NAVY, bold=True, caps=True)
    p = cell.add_paragraph()
    set_paragraph(p, before=0, after=0, line=1.12)
    r = p.add_run(text)
    set_run(r, size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def para(doc, text="", size=10.5, color=BLACK, bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after, line=1.12, align=align)
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p


def rich_para(doc, parts, before=0, after=6, line=1.12, align=None):
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after, line=line, align=align)
    for text, opts in parts:
        r = p.add_run(text)
        set_run(
            r,
            font=opts.get("font", "Aptos"),
            size=opts.get("size", 10.5),
            color=opts.get("color", BLACK),
            bold=opts.get("bold"),
            italic=opts.get("italic"),
            caps=opts.get("caps", False),
        )
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run(r, size=10.3, color=BLACK)
    return p


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_matrix_table(doc, headers, rows, widths, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_layout(table, widths)
    set_table_borders(table, color=BORDER, size="4")
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        format_cell(hdr[i], size=8.8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_shading(cells[i], WHITE if len(table.rows) % 2 == 0 else "FBFCFE")
            format_cell(cells[i], size=8.9, color=BLACK)
    set_table_layout(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def doc_separator(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_internal_memo():
    doc = setup_doc("16Arena Final Internal Legal Determination Memorandum")
    running_header_footer(doc, "16Arena | Internal Legal Determination", "16Arena - Internal Legal Determination")
    add_masthead(
        doc,
        "Privileged and Confidential - Attorney Work Product",
        "Internal Legal Determination Memorandum",
        "PROG Act, 2025 and PROG Rules, 2026 - classification of 16Arena's free-entry tournaments",
        [
            ("Memo No.", "16A/ILD/2026/FINAL"),
            ("Date", "[insert date], May 2026"),
            ("From", "General Counsel / Authorised Director, 16Arena"),
            ("To", "Board of Directors, 16Arena"),
            ("Subject", "Determination that the free-entry tournament feature, as architected, should not be treated as an online money game or online money gaming service"),
            ("Core Sources", "Promotion and Regulation of Online Gaming Act, 2025; G.S.R. 303(E) dated 22 April 2026; S.O. 1992(E) and S.O. 1994(E) dated 22 April 2026; RBI PPI Master Directions, 2021; Income-tax provisions on online gaming winnings"),
        ],
    )

    add_callout(
        doc,
        "Board conclusion",
        "Subject to the hard implementation conditions recorded in this memorandum, 16Arena's free-entry cash-prize tournaments should be treated as outside the definition of an online money game because the user pays no fee, makes no deposit, and places no money or other stake for participation. This conclusion must remain tied to the actual product architecture, payment flow, tax setup, KYC process, state restrictions, and change-control commitments set out below.",
        fill=LIGHT_GOLD,
    )

    heading(doc, "1. Purpose and Controlled Use", 1)
    para(doc, "1.1 This memorandum records the internal legal determination of 16Arena in relation to its proposed free-entry cash-prize tournament feature and the separated Game Shop feature.")
    para(doc, "1.2 It is intended for adoption by the Board, for controlled disclosure to the partner bank, the RBI-authorised payment partner, statutory advisors, and the Online Gaming Authority of India, if and when an application or representation is filed.")
    para(doc, "1.3 This memorandum does not by itself authorise launch. Launch is conditional on the legal, payments, tax, accounting, secretarial, data protection, state-law and operational controls recorded in this document.")

    heading(doc, "2. Current Legal Framework", 1)
    para(doc, "2.1 The Promotion and Regulation of Online Gaming Act, 2025 came into force on 1 May 2026 pursuant to S.O. 1994(E) dated 22 April 2026. The Online Gaming Authority of India was constituted under S.O. 1992(E) dated 22 April 2026. The Promotion and Regulation of Online Gaming Rules, 2026 were notified under G.S.R. 303(E) dated 22 April 2026 and also came into force on 1 May 2026.")
    para(doc, "2.2 Section 2(1)(g) defines an online money game as an online game played by a user by paying fees, depositing money or other stakes in expectation of winning monetary or other enrichment in return for such money or other stakes. Section 2(1)(h) defines online money gaming service as a service offered for entering or playing an online money game.")
    para(doc, "2.3 Section 5 prohibits offering an online money game or online money gaming service. Section 7 prohibits banks, financial institutions and payment facilitators from facilitating transactions towards payment for any online money gaming service. Sections 9 to 11 create material penal and officer-liability exposure.")
    para(doc, "2.4 Rules 8, 9 and 10 prescribe the determination framework. Rule 8 states when determination is required; Rule 9 sets out the factors for determination; Rule 10 sets out the procedure and effect of a determination order. Rule 12 concerns requirement for registration, and should not be described as the general determination provision.")

    heading(doc, "3. Platform Features Reviewed", 1)
    add_matrix_table(
        doc,
        ["Feature", "Architecture recorded by the Board", "Compliance treatment"],
        [
            ("Feature A: Free-entry tournaments", "Users participate without paying any entry fee, making any deposit, purchasing any token, holding any paid credit, or staking any value. Winners may receive company-funded cash prizes.", "Treat as not an online money game, subject to the controls in Sections 5 to 8 below."),
            ("Prize payout ledger", "Amounts payable to winners are recorded as a company liability in an internal payout ledger. It is not user-loadable, not transferable, not usable for purchases, and not described as a cash wallet.", "Use 'internal prize ledger' or 'payout payable'. Avoid 'wallet' terminology unless payments counsel confirms no PPI issue."),
            ("Feature B: Game Shop", "A separated checkout/stored-value feature for purchasing gaming vouchers or digital goods. It must not fund tournament entry and prize money must not flow into it.", "Do not launch or expand unless payments counsel confirms whether it is a closed-system arrangement or requires a licensed PPI/PA structure."),
        ],
        [1900, 3800, 3660],
    )

    heading(doc, "4. Determination on Free-Entry Tournaments", 1)
    para(doc, "4.1 On the facts recorded, the statutory trigger for an online money game is absent. The user does not pay a fee, does not deposit money, and does not place any money or other stake for participation.")
    para(doc, "4.2 The fact that a winner may receive a company-funded prize should not, by itself, convert a free-entry tournament into an online money game. The essential link required by Section 2(1)(g) is participation by paying fees, depositing money or other stakes in expectation of enrichment in return for such money or other stakes.")
    para(doc, "4.3 The Board records that the following statements must remain true for every cash-prize tournament:")
    for item in [
        "No entry fee, registration fee, participation fee, subscription, battle pass, premium membership or access fee is charged for tournament participation.",
        "No user deposit is accepted into any tournament account, tournament wallet or tournament ledger.",
        "No coin, token, credit, voucher, object or in-game item purchased for money is required, consumed or blocked for tournament participation.",
        "Prize money is not funded from user deposits or pooled user contributions.",
        "The Game Shop feature, any cashback, and any promotional credit cannot be used to enter or enhance cash-prize tournaments.",
        "Cash prizes are disbursed only after KYC, TDS processing and payout approval through a lawful banking/payment arrangement.",
    ]:
        bullet(doc, item)

    heading(doc, "5. Payment Architecture and PPI Risk Controls", 1)
    add_callout(
        doc,
        "Hard payment condition",
        "The earlier legal opinion identified a serious PPI/payment-system risk if a rupee-denominated, withdrawable in-app cash wallet is operated by 16Arena without RBI authorisation. The Board therefore records that the prize payout balance shall be implemented only as an internal accounting ledger and payout payable, and actual money movement shall be routed through a bank/payment partner structure cleared by qualified payments counsel.",
        fill="FFF7ED",
        border=GOLD,
    )
    for item in [
        "No user-loaded cash wallet shall be operated for the tournament module unless 16Arena obtains the required authorisation or uses an authorised issuer structure.",
        "No prize balance shall be transferable between users or usable to purchase goods, vouchers, subscriptions or in-app items.",
        "The prize ledger shall be reconciled daily with a segregated prize-funding account or other account structure approved by the CA and payments counsel.",
        "The partner bank and payment aggregator shall receive the final product flow, payout SOP, KYC SOP, tax note and this memorandum before launch.",
    ]:
        bullet(doc, item)

    heading(doc, "6. OGAI Position and External Representation", 1)
    para(doc, "6.1 Under the final 2026 Rules, determination is not automatically mandatory for every online game. However, given the presence of cash prizes and financial flows, the Board authorises filing an application-cum-representation with the Online Gaming Authority of India seeking a determination or clarification that Feature A is not an online money game.")
    para(doc, "6.2 16Arena shall not market, represent or display the platform as determined, approved, recognised or registered by OGAI unless an actual determination order or certificate has been issued and remains valid.")
    para(doc, "6.3 If a determination order is issued, any change in payment facilitation, access model, tournament entry conditions, prize funding, Game Shop linkage or monetisation model must be notified or freshly reviewed before deployment.")

    heading(doc, "7. Compliance Preconditions", 1)
    add_matrix_table(
        doc,
        ["Area", "Required control before launch", "Owner"],
        [
            ("Tax/TDS", "TAN, TDS deduction logic for net winnings, 26Q/Form 16A process, CA-approved payout accounting.", "CA / Accounts"),
            ("GST", "Written CA note confirming GST treatment of free-entry prize disbursal and separate treatment of advertising, subscriptions, in-app purchases or Game Shop revenue.", "CA"),
            ("Company law", "Board approval, authorised signatory, MOA/object clause review, register and filing check.", "CS"),
            ("Payments", "RBI-authorised PA/payout arrangement or payments counsel note confirming lawful routing.", "Payments counsel / Finance"),
            ("KYC", "No manual storage of Aadhaar unless lawfully permitted; use licensed provider or lawful OVD/PAN process with consent and retention controls.", "Compliance / Product"),
            ("Data protection", "DPDP-compliant notice, consent logs, access controls, encryption, breach procedure and retention schedule.", "DPO / Tech"),
            ("State restrictions", "Geo-restrict states identified by counsel; keep state-law matrix current.", "Legal / Product"),
            ("Game classification", "Game-by-game skill and regulatory review before cash-prize launch.", "Legal"),
            ("User safety", "18+ age gate, grievance redress, reporting controls, fair-play and integrity monitoring.", "Compliance / Product"),
        ],
        [1700, 5600, 2060],
    )

    heading(doc, "8. Board Commitments", 1)
    for item in [
        "The Board adopts the no-fee, no-deposit, no-stake architecture as a mandatory product constraint.",
        "The Board authorises management to file an OGAI application-cum-representation in the corrected form approved by counsel.",
        "The Board directs management to obtain CA and CS domain notes before launch and to maintain them with this memorandum.",
        "The Board requires fresh legal sign-off for any feature that introduces user-side payment, deposit, stake, subscription gating, paid token use, prize-to-shop transfer, or wallet cash-out.",
    ]:
        bullet(doc, item)

    heading(doc, "9. Caveats and Limitations", 1)
    para(doc, "9.1 This memorandum is based on the facts and architecture recorded above. Any factual departure may change the analysis.")
    para(doc, "9.2 This memorandum does not constitute a payments authorisation opinion, GST opinion, accounting opinion, FEMA opinion, or state-specific gaming-law opinion. Those matters require separate domain review.")
    para(doc, "9.3 This memorandum is privileged and confidential and may be disclosed externally only on a need-to-know basis with approval of the authorised director or legal function.")

    para(doc, "For and on behalf of 16Arena", before=16, after=18, bold=True, color=NAVY)
    para(doc, "__________________________________", after=2)
    para(doc, "Name: [insert authorised signatory]", after=2)
    para(doc, "Designation: [insert designation]", after=2)
    para(doc, "Date / Place: [insert]", after=2)

    path = OUT / "16Arena_Final_Internal_Legal_Determination_Memo.docx"
    doc.save(path)
    return path


def build_ogai_application():
    doc = setup_doc("16Arena Final OGAI Petition for Determination and Permission", premium=True)
    running_header_footer(doc, "16Arena | Petition before OGAI", "Petition for Determination and Permission", "Privileged and Confidential")

    p = doc.add_paragraph()
    set_paragraph(p, before=2, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("BEFORE THE ONLINE GAMING AUTHORITY OF INDIA")
    set_run(r, font="Georgia", size=13, color=NAVY, bold=True, caps=True)
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("(Constituted under Section 8 of the Promotion and Regulation of Online Gaming Act, 2025)")
    set_run(r, font="Georgia", size=9.5, color=MUTED, italic=True)

    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("Diary / Application No.: ____________________")
    set_run(r, font="Georgia", size=10, color=BLACK)

    table = doc.add_table(rows=1, cols=1)
    set_table_layout(table, [9360])
    set_table_borders(table, color=GOLD, size="8")
    set_cell_shading(table.rows[0].cells[0], LIGHT_GOLD)
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, before=3, after=3, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("PETITION / APPLICATION")
    set_run(r, font="Georgia", size=18, color=NAVY, bold=True, caps=True)
    p = cell.add_paragraph()
    set_paragraph(p, before=0, after=2, line=1.08, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("Seeking determination, permission and such other appropriate order in respect of 16Arena's free-entry cash-prize tournament feature")
    set_run(r, font="Georgia", size=11, color=SLATE, bold=True)
    p = cell.add_paragraph()
    set_paragraph(p, before=0, after=4, line=1.08, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("Under Section 8(2)(a) of the PROG Act, 2025 read with Rules 8, 9, 10 and 23 of the PROG Rules, 2026")
    set_run(r, font="Georgia", size=9.5, color=GOLD_DARK, bold=True)

    heading(doc, "In the Matter of", 2)
    rich_para(
        doc,
        [
            ("The Promotion and Regulation of Online Gaming Act, 2025; ", {"font": "Georgia", "size": 10.5}),
            ("and", {"font": "Georgia", "size": 10.5, "italic": True}),
            (" the Promotion and Regulation of Online Gaming Rules, 2026.", {"font": "Georgia", "size": 10.5}),
        ],
        after=3,
    )
    heading(doc, "And in the Matter of", 2)
    para(doc, "A petition/application seeking determination that the free-entry tournament feature of the online platform 16Arena is not an online money game or online money gaming service, and seeking permission/approval to operationalise the said feature subject to the compliance undertakings recorded herein.", size=10.5, after=8)

    heading(doc, "And in the Matter of", 2)
    parties = doc.add_table(rows=2, cols=2)
    set_table_layout(parties, [6850, 2510])
    set_table_borders(parties, color=BORDER, size="4")
    parties.rows[0].cells[0].text = "[insert legal name of Applicant Company]\nA company incorporated under the Companies Act, 2013\nCIN: [insert CIN]\nRegistered Office: [insert registered office]\nThrough: [insert authorised representative]"
    parties.rows[0].cells[1].text = "... Applicant / Petitioner"
    parties.rows[1].cells[0].text = "The Online Gaming Authority of India\nMinistry of Electronics and Information Technology\nGovernment of India"
    parties.rows[1].cells[1].text = "... Authority"
    for row in parties.rows:
        for cell in row.cells:
            set_cell_shading(cell, "FBFCFE")
            format_cell(cell, font="Georgia", size=9.5, color=BLACK)

    p = doc.add_paragraph()
    set_paragraph(p, before=10, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("PETITION / APPLICATION ON BEHALF OF THE APPLICANT")
    set_run(r, font="Georgia", size=12, color=NAVY, bold=True, caps=True)
    paragraph_border_bottom(p, color=GOLD, size="8", space="2")

    heading(doc, "Index of Filing", 1)
    add_matrix_table(
        doc,
        ["Sl.", "Particulars", "Page No."],
        [
            ("1", "Synopsis and List of Dates", "____"),
            ("2", "Memo of Petition / Application", "____"),
            ("3", "Grounds for determination and permission", "____"),
            ("4", "Undertakings and compliance conditions", "____"),
            ("5", "Prayer", "____"),
            ("6", "Verification", "____"),
            ("7", "Annexures A-1 to A-12", "____"),
        ],
        [700, 7400, 1260],
        header_fill=NAVY,
    )

    heading(doc, "Synopsis and List of Dates", 1)
    para(doc, "The present petition/application is being filed by the Applicant, the operator of the online platform 16Arena, seeking an authoritative determination and necessary permission/approval from the Online Gaming Authority of India in respect of its proposed free-entry cash-prize tournament feature.", size=10.5)
    para(doc, "The Applicant's case is that the said feature does not fall within the definition of an online money game under Section 2(1)(g) of the Promotion and Regulation of Online Gaming Act, 2025, because a user is not required to pay any entry fee, make any deposit, purchase any token, acquire any paid access, or stake any value for participating in the tournament.", size=10.5)
    para(doc, "The Applicant respectfully submits that the mere disbursal of company-funded prize money to a successful participant, after KYC, TDS and compliance checks, does not convert the feature into an online money game where the user has not put in money or other stakes. The prize payout is not a return on user consideration; it is a company-funded reward under a free-entry promotional/tournament structure.", size=10.5)
    para(doc, "The Applicant has, therefore, approached this Hon'ble Authority in advance of operationalisation, in a transparent and bona fide manner, so that its proposed structure may be examined against the Act and the Rules, and so that partner banks, payment aggregators, statutory advisors and users may proceed on the basis of clear regulatory comfort.", size=10.5)
    para(doc, "The Applicant further undertakes that the free-entry tournament feature shall remain fully separated from any paid Game Shop or digital-commerce feature, that no paid wallet or paid credit shall fund tournament participation, that the internal prize ledger shall not operate as a user cash wallet, and that any material change in the architecture shall be placed for fresh legal and regulatory review.", size=10.5)
    para(doc, "The present petition/application is accordingly filed seeking a determination, no-objection, permission, approval, or such other written regulatory communication as this Hon'ble Authority may deem appropriate, permitting the Applicant to operationalise the free-entry tournament feature subject to the undertakings and compliance conditions set out herein.", size=10.5, after=10)

    add_matrix_table(
        doc,
        ["Date / Stage", "Event"],
        [
            ("[insert date]", "The Applicant was incorporated under the Companies Act, 2013 and commenced preparatory work for operating the 16Arena platform."),
            ("[insert period]", "The Applicant developed the 16Arena platform architecture, including a proposed free-entry tournament feature and a separate Game Shop / digital-commerce feature."),
            ("22 August 2025", "The Promotion and Regulation of Online Gaming Act, 2025 received Presidential assent."),
            ("22 April 2026", "The Central Government constituted the Online Gaming Authority of India under Section 8 of the Act."),
            ("22 April 2026", "The Promotion and Regulation of Online Gaming Rules, 2026 were notified. The Rules prescribe, among other things, the factors for determination and the information to be disclosed at the time of application."),
            ("1 May 2026", "The Act and the 2026 Rules came into force."),
            ("[insert date]", "The Applicant finalised the proposed no-fee, no-deposit and no-stake tournament architecture for regulatory review."),
            ("20 May 2026", "The Applicant obtained legal opinion on the proposed 16Arena structure and compliance conditions, including payment, tax, KYC, data protection and gaming-law issues."),
            ("[insert date]", "The Applicant prepared an internal legal determination memorandum recording the product architecture, compliance assumptions and launch preconditions."),
            ("[insert date]", "The Applicant initiated review by its Chartered Accountant for TDS, GST and accounting treatment and by its Company Secretary for board approvals, filings, objects clause and statutory compliances."),
            ("[insert date]", "The Applicant initiated engagement with partner bank/payment aggregator for lawful payout routing and banking diligence."),
            ("[insert date]", "The Board of the Applicant approved filing of this petition/application and authorised the signatory."),
            ("[insert date]", "The present petition/application is filed before this Hon'ble Authority seeking determination and necessary permission/approval/no-objection for operationalising the free-entry tournament feature."),
        ],
        [1800, 7560],
        header_fill=TEAL,
    )

    para(doc, "", after=8)
    p = doc.add_paragraph()
    set_paragraph(p, before=2, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("BEFORE THE ONLINE GAMING AUTHORITY OF INDIA")
    set_run(r, font="Georgia", size=12, color=NAVY, bold=True, caps=True)
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("PETITION / APPLICATION UNDER SECTION 8(2)(a) OF THE PROG ACT, 2025 READ WITH RULES 8, 9, 10 AND 23 OF THE PROG RULES, 2026")
    set_run(r, font="Georgia", size=10.5, color=GOLD_DARK, bold=True, caps=True)

    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=6, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("APPLICATION")
    set_run(r, font="Georgia", size=16, color=NAVY, bold=True, caps=True)
    paragraph_border_bottom(p, color=GOLD, size="8", space="2")

    rich_para(
        doc,
        [("MOST RESPECTFULLY SHOWETH:", {"font": "Georgia", "size": 11, "bold": True, "color": NAVY, "caps": True})],
        before=4,
        after=8,
    )

    heading(doc, "1. Jurisdiction, Maintainability and Nature of Relief", 1)
    para(doc, "1.1 The Authority is empowered under Section 8(2)(a) of the Promotion and Regulation of Online Gaming Act, 2025 to determine, on receipt of an application or suo motu, whether a particular online game is an online money game or otherwise.", size=10.5)
    para(doc, "1.2 Rules 8, 9 and 10 of the Promotion and Regulation of Online Gaming Rules, 2026 prescribe the manner, objective factors and procedure for determination. Rule 23 prescribes the information to be disclosed at the time of application for determination or registration, including information relating to user safety features and grievance redressal.", size=10.5)
    para(doc, "1.3 The present petition/application is filed voluntarily and in abundant caution. The Applicant's respectful position is that no mandatory determination is presently triggered under Rule 8(1) unless the Authority directs determination, the game is intended to be offered as an e-sport, or the Central Government notifies a relevant category. Nevertheless, since the feature involves company-funded cash prizes and will necessarily require bank/payment partner diligence, the Applicant seeks an authoritative order, no-objection, permission or written determination from this Hon'ble Authority before operationalisation.", size=10.5)
    para(doc, "1.4 The Applicant does not seek permission to conduct any online money game. The permission/approval sought is for operationalising the free-entry tournament feature upon a determination that the said feature, as described and undertaken herein, is not an online money game or online money gaming service.", size=10.5)

    heading(doc, "2. Particulars of the Applicant", 1)
    add_matrix_table(
        doc,
        ["Sl.", "Particular", "Details"],
        [
            ("2.1", "Legal name", "[insert legal name]"),
            ("2.2", "Type of entity", "Private limited company / [insert]"),
            ("2.3", "Date of incorporation", "[insert]"),
            ("2.4", "Corporate Identity Number", "[insert CIN]"),
            ("2.5", "PAN / TAN", "[insert PAN] / [insert TAN if obtained]"),
            ("2.6", "Registered office", "[insert address]"),
            ("2.7", "Principal place of business", "[insert address]"),
            ("2.8", "Authorised representative", "[insert name, designation, email and phone]"),
            ("2.9", "Platform name", "16Arena"),
            ("2.10", "Mode of access", "Mobile application and web interface"),
            ("2.11", "Launch status", "Proposed / [insert actual status]"),
        ],
        [700, 2600, 6060],
        header_fill=NAVY,
    )

    heading(doc, "3. Facts Giving Rise to the Present Petition", 1)
    para(doc, "3.1 The Applicant operates, or proposes to operate, the online platform 16Arena. The platform is intended to host esports/casual gaming tournaments and associated digital-commerce features, subject to applicable laws and regulatory directions.", size=10.5)
    para(doc, "3.2 The subject matter of the present petition is strictly limited to the free-entry cash-prize tournament feature. Under this feature, a user does not pay any entry fee, platform fee, registration fee, contest fee, subscription charge or other consideration for participation.", size=10.5)
    para(doc, "3.3 A user is also not required to deposit money into any account, purchase or hold any token, buy any coin, acquire any paid access product, expend any virtual object, subscribe to any membership, buy any voucher, or stake any asset of monetary or equivalent value for entering or playing the tournament.", size=10.5)
    para(doc, "3.4 Winners may receive company-funded cash prizes. Such prizes are not funded from user entry fees or user deposits. The prize amount, until payout, is recorded only as an internal payout liability/prize ledger of the Applicant. The said internal record is not a user-loadable wallet, not transferable between users, not usable for purchases and not available for funding any tournament entry.", size=10.5)
    para(doc, "3.5 Payout is proposed to be made only to the user's KYC-verified bank account or UPI handle after tax deduction, compliance checks, internal approval and lawful banking/payment processing. The Applicant does not seek to operate any unauthorised payment system or withdrawable user cash wallet.", size=10.5)
    para(doc, "3.6 The platform may separately contain a Game Shop / digital-commerce feature for purchase of gaming vouchers or digital goods. The Applicant undertakes that this feature shall remain technically and financially firewalled from tournaments. Game Shop balance, cashback, vouchers, credits or other paid value shall not be used directly or indirectly for tournament participation, ranking, eligibility, prize enhancement or payout.", size=10.5)
    para(doc, "3.7 The Applicant is placing this structure before the Authority at the pre-operational / pre-scale stage so that the Authority may examine the architecture and issue an appropriate determination, no-objection, permission, approval, clarification or direction before the Applicant proceeds with full launch.", size=10.5)

    heading(doc, "4. Architecture Placed Before the Authority", 1)
    add_matrix_table(
        doc,
        ["Component", "Architecture", "Regulatory relevance"],
        [
            ("Free-entry tournaments", "Users participate without paying any fee, making any deposit, purchasing or using any token, or staking any value.", "No user-side money or stake is involved."),
            ("Prize payout ledger", "Internal accounting record of company liability. It is not loadable, transferable, spendable or usable for purchases.", "Avoids characterisation as a stake or spendable cash wallet."),
            ("Payout flow", "Payout only to KYC-verified bank account or UPI handle after TDS and compliance checks, through lawful banking/payment rails.", "Supports auditable prize disbursal."),
            ("Game Shop", "Separated digital-commerce feature. No value flows from Game Shop to tournaments and no prize money flows to Game Shop.", "Prevents indirect stake creation."),
        ],
        [2100, 4650, 2710],
        header_fill=TEAL,
    )

    heading(doc, "5. Compliance Undertakings", 1)
    for item in [
        "The Applicant shall not introduce any entry fee, deposit, user stake, subscription gate, battle pass, paid token or paid in-game asset requirement for the free-entry tournament feature without prior legal review and, where required, fresh engagement with the Authority.",
        "The Applicant shall not describe or operate the internal prize ledger as a user cash wallet and shall not permit it to be loaded, transferred, spent, used for purchases, or linked to Game Shop balances.",
        "The Applicant shall route payouts only through lawful banking/payment partner arrangements after KYC, TDS and compliance checks.",
        "The Applicant shall maintain age-gating, state restrictions, grievance redressal, fair-play/integrity monitoring, data protection controls, and such user-safety features as may be applicable.",
        "The Applicant shall comply with all lawful directions, orders, guidelines and codes of practice issued by the Central Government or the Authority.",
        "The Applicant shall not represent the platform as determined, approved, permitted or registered by the Authority unless an order/certificate/communication to that effect has actually been issued and remains valid.",
    ]:
        bullet(doc, item)

    p = doc.add_paragraph()
    set_paragraph(p, before=14, after=6, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("GROUNDS")
    set_run(r, font="Georgia", size=16, color=NAVY, bold=True, caps=True)
    paragraph_border_bottom(p, color=GOLD, size="8", space="2")

    grounds = [
        ("A.", "Because Section 2(1)(g) of the Act requires a user to play an online game by paying fees, depositing money or other stakes in expectation of winning monetary or other enrichment in return for such money or other stakes."),
        ("B.", "Because the Applicant's free-entry tournament feature involves no user-side fee, no deposit of money, no paid access, no subscription gate, no paid token, and no other stake at any stage of participation."),
        ("C.", "Because company-funded prize disbursal to winners is not the same as enrichment in return for money or stakes deposited by the user."),
        ("D.", "Because Rule 9 factors, when applied to the Applicant's architecture, support a finding that the free-entry tournament feature is not an online money game."),
        ("E.", "Because the internal prize ledger is a non-transferable payout liability record and is not a cash wallet, PPI, stake wallet or user-spendable stored-value instrument for tournament participation."),
        ("F.", "Because the Game Shop feature is technically and financially separated from the tournament feature and cannot be used directly or indirectly to participate in cash-prize tournaments."),
        ("G.", "Because an authoritative determination/permission from the Authority would enable compliant launch, bank/payment partner diligence, user safety and transparent regulatory supervision."),
        ("H.", "Because the prohibitions under Sections 5 and 7 are directed at online money games and online money gaming services, and should not be attracted to a free-entry tournament payout flow that does not involve payment by the user for entering or playing the game."),
        ("I.", "Because the Applicant has approached the Authority voluntarily and before full operationalisation, demonstrating bona fides, transparency and willingness to be regulated in accordance with the Act, Rules, guidelines, codes of practice and directions issued by the Central Government or the Authority."),
        ("J.", "Because regulatory clarity at this stage will prevent avoidable disruption by banks, payment aggregators and service providers, while allowing the Applicant to implement age-gating, state restrictions, KYC, TDS, AML, grievance redressal, user-safety and data-protection controls from inception."),
        ("K.", "Because the Applicant undertakes not to market, represent or advertise the platform as approved, permitted, determined or registered by the Authority unless and until an actual order, certificate or written communication to that effect is issued and remains valid."),
        ("L.", "Because the Applicant undertakes to seek fresh legal review and, where required, fresh regulatory engagement before introducing any feature involving user-side payment, deposit, stake, paid token, subscription gate, paid Game Shop linkage, prize-to-shop transfer or any materially altered monetisation model."),
    ]
    for label, text in grounds:
        rich_para(doc, [(label + " ", {"font": "Georgia", "size": 10.5, "bold": True, "color": NAVY}), (text, {"font": "Georgia", "size": 10.5, "color": BLACK})], after=6)

    p = doc.add_paragraph()
    set_paragraph(p, before=14, after=6, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("PRAYERS")
    set_run(r, font="Georgia", size=16, color=NAVY, bold=True, caps=True)
    paragraph_border_bottom(p, color=GOLD, size="8", space="2")

    para(doc, "In view of the facts, grounds and undertakings stated above, the Applicant most respectfully prays that this Hon'ble Authority may be pleased to:", size=10.5)
    prayers = [
        "take the present petition/application on record;",
        "call for, receive and examine such additional information, demonstration, product-flow materials, technical architecture documents, payment-flow materials and compliance affidavits as this Hon'ble Authority may consider necessary;",
        "grant the Applicant an opportunity of hearing, presentation or product demonstration before passing final orders, if the Authority considers such hearing useful;",
        "determine under Section 8(2)(a) of the Act read with Rules 8, 9 and 10 of the Rules that the free-entry cash-prize tournament feature of 16Arena, as described in this petition/application, is not an online money game;",
        "record that the Applicant is not offering an online money gaming service in respect of the said feature as described herein;",
        "grant the necessary permission, approval, no-objection, determination order or written regulatory communication, as may be considered appropriate, permitting the Applicant to operationalise the said free-entry tournament feature subject to the undertakings and compliance conditions recorded herein;",
        "record that partner banks, payment aggregators and service providers facilitating lawful prize disbursal under the architecture described herein are not, merely by facilitating such payout, facilitating a transaction towards payment for an online money gaming service within Section 7 of the Act;",
        "clarify that registration is not required unless the Central Government notifies the relevant category under Rule 12 or unless any game is intended to be offered as an e-sport;",
        "take on record the undertakings contained in this petition/application and permit the Applicant to file updated undertakings if required by the Authority;",
        "permit the Applicant to proceed with partner bank, payment aggregator, tax, secretarial, KYC, AML, user-safety and data-protection arrangements required for launch while the determination process is pending, without representing final approval until the Authority issues its order or communication;",
        "direct that any determination, permission, approval, no-objection or written communication issued in favour of the Applicant shall remain subject to the Applicant maintaining the no-fee, no-deposit, no-stake architecture described herein;",
        "permit the Applicant to approach the Authority for clarification or modification if there is any future change in product architecture, monetisation model, payment flow or regulatory position; and",
        "pass such further or other order(s) as this Hon'ble Authority may deem fit in the facts and circumstances of the case.",
    ]
    for idx, item in enumerate(prayers, 1):
        rich_para(doc, [(f"{idx}. ", {"font": "Georgia", "size": 10.5, "bold": True, "color": NAVY}), (item, {"font": "Georgia", "size": 10.5, "color": BLACK})], after=4)

    heading(doc, "8. Interim / Ad-Interim Request", 1)
    para(doc, "Pending final determination, the Applicant respectfully requests that the Authority may acknowledge receipt of the present filing and permit the Applicant to engage with partner banks, payment aggregators and statutory advisors on the basis that the Applicant has voluntarily approached the Authority and has undertaken not to represent any final approval until the Authority issues its order or communication.", size=10.5)

    heading(doc, "9. List of Annexures", 1)
    add_matrix_table(
        doc,
        ["Annexure", "Document"],
        [
            ("A-1", "Certificate of Incorporation, PAN, TAN and CIN particulars of the Applicant"),
            ("A-2", "Board resolution authorising filing and appointing authorised representative"),
            ("A-3", "Final Internal Legal Determination Memorandum dated [insert] May 2026"),
            ("A-4", "Legal Opinion of Shashwat, Advocate, dated 20 May 2026"),
            ("A-5", "Functional product walkthrough and screenshots"),
            ("A-6", "Architecture diagram showing separation between free-entry tournaments and Game Shop"),
            ("A-7", "Payment flow, PA/bank arrangement note and payout SOP"),
            ("A-8", "CA note on TDS, GST and accounting treatment"),
            ("A-9", "CS note on company-law approvals, objects clause and compliance calendar"),
            ("A-10", "Terms of Service, Tournament Rules, Privacy Notice and Grievance Policy"),
            ("A-11", "KYC, AML, age-gating and state-restriction SOPs"),
            ("A-12", "Any other document directed or requested by the Authority"),
        ],
        [1200, 8160],
        header_fill=NAVY,
    )

    heading(doc, "10. Verification", 1)
    para(doc, "I, [insert name], [insert designation] of [insert legal name of Applicant], duly authorised by the Board of Directors by resolution dated [insert], do hereby verify that the contents of paragraphs 1 to 10 of the present petition/application are true and correct to my knowledge and belief, based on the records of the Applicant and information believed to be true. No material fact has been concealed.", size=10.5)
    para(doc, "Verified at [insert place] on this [insert] day of May 2026.", size=10.5, after=18)
    para(doc, "For [insert legal name of Applicant]", bold=True, color=NAVY, after=18)
    para(doc, "__________________________________", after=2)
    para(doc, "Authorised Representative", after=2)
    para(doc, "Name / Designation: [insert]", after=2)
    para(doc, "Email / Phone: [insert]", after=2)

    path = OUT / "16Arena_Final_OGAI_Determination_Application.docx"
    doc.save(path)
    return path


def build_counsel_representation():
    doc = setup_doc("16Arena Final Counsel Representation", premium=True)
    running_header_footer(doc, "Shashwat | Counsel Representation", "Counsel's Representation - 16Arena", "Privileged | Addressees Only")
    add_letterhead(doc)

    meta = doc.add_table(rows=3, cols=2)
    set_table_layout(meta, [1450, 7910])
    set_table_borders(meta, color="E5E7EB", size="4")
    for i, (k, v) in enumerate([
        ("Ref.", "SH/16A/2026/FINAL"),
        ("Date", "[insert date], May 2026"),
        ("Matter", "Engagement with 16Arena - legal comfort for bank and payment aggregator diligence"),
    ]):
        c0, c1 = meta.rows[i].cells
        c0.text = k
        c1.text = v
        set_cell_shading(c0, NAVY if i == 0 else LIGHT_BLUE)
        set_cell_shading(c1, WHITE)
        format_cell(c0, size=8.8, color=WHITE if i == 0 else NAVY, bold=True)
        format_cell(c1, size=9.2, color=BLACK)

    para(doc, "To,", before=10, after=2)
    para(doc, "The Compliance Officer", after=1, bold=True)
    para(doc, "[Name of Partner Bank]", after=1)
    para(doc, "[Address]", after=4)
    para(doc, "The Compliance Officer", after=1, bold=True)
    para(doc, "[Name of Payment Aggregator]", after=1)
    para(doc, "[Address]", after=10)

    rich_para(
        doc,
        [
            ("Re: ", {"bold": True, "color": NAVY}),
            ("16Arena - representation on the non-application of Section 7 of the Promotion and Regulation of Online Gaming Act, 2025 to the free-entry tournament payout flow", {"bold": True, "color": NAVY}),
        ],
        after=8,
    )

    add_callout(
        doc,
        "Executive comfort",
        "On the facts and documents represented to me, 16Arena's free-entry tournament feature does not involve a user-side fee, deposit of money or other stake. The prize payout flow is therefore not a transaction towards payment for an online money gaming service. Section 7 of the PROG Act should not be engaged by the bank or payment aggregator merely facilitating the lawful payout of company-funded prizes, subject always to the controls and limitations set out in this representation.",
        fill=LIGHT_GOLD,
        border=GOLD,
    )

    para(doc, "Dear Sirs / Madam,", after=8)

    heading(doc, "1. Instructions and Documents Reviewed", 1)
    para(doc, "1.1 I have been requested by the operator of the online platform 16Arena to provide this representation for the limited assistance of the addressee bank and payment aggregator in their diligence regarding 16Arena's proposed free-entry tournament payout flow.")
    para(doc, "1.2 I have proceeded on the basis of the product description, internal legal determination memorandum, payout flow, Game Shop separation description, draft user terms, and the legal opinion dated 20 May 2026. I have not independently audited source code, bank systems, tax filings, KYC records, or live product implementation.")

    heading(doc, "2. Question Presented", 1)
    para(doc, "2.1 The question is whether the addressee bank or payment aggregator, by facilitating payouts of company-funded prizes to winners of 16Arena's free-entry tournaments, would be facilitating a transaction towards payment for an online money gaming service within Section 7 of the Promotion and Regulation of Online Gaming Act, 2025.")

    heading(doc, "3. Applicable Statutory Position", 1)
    para(doc, "3.1 Section 2(1)(g) defines an online money game by reference to an online game played by a user by paying fees, depositing money or other stakes in expectation of winning monetary or other enrichment in return for such money or other stakes.")
    para(doc, "3.2 Section 7 prohibits banks, financial institutions and other payment facilitators from facilitating transactions towards payment for any online money gaming service. The prohibition is therefore engaged where the payment is towards entering or playing an online money game.")
    para(doc, "3.3 The final Promotion and Regulation of Online Gaming Rules, 2026, effective 1 May 2026, use Rules 8, 9 and 10 for determination. Rule 9 requires consideration of user-side payments, deposits or stakes, expectation of winnings in return for such money or stakes, the revenue model, and the manner in which rewards or assets are redeemed or monetised.")

    heading(doc, "4. Facts Assumed for this Representation", 1)
    add_matrix_table(
        doc,
        ["Assumed fact", "Compliance significance"],
        [
            ("Users pay no entry fee, deposit no money, and place no stake to enter Feature A tournaments.", "The statutory trigger under Section 2(1)(g) is absent."),
            ("Prize money is funded by 16Arena and not by user contributions or pooled deposits.", "Payout is a company-funded prize disbursal, not user-stake settlement."),
            ("The prize ledger is internal, non-loadable, non-transferable and not usable for purchases.", "Reduces PPI and stake-characterisation risk."),
            ("The Game Shop feature is firewalled and cannot fund or enhance tournament entry.", "Prevents indirect stake creation through a paid feature."),
            ("Payouts occur only to KYC-verified bank/UPI details after tax and compliance checks.", "Supports lawful payout and auditability."),
        ],
        [3850, 5510],
        header_fill=NAVY,
    )

    heading(doc, "5. Opinion", 1)
    para(doc, "5.1 On the facts assumed, 16Arena's free-entry tournament feature should not be treated as an online money game because there is no user-side fee, deposit of money or other stake. The user may expect a prize, but not in return for money or other stake contributed by the user.")
    para(doc, "5.2 It follows that 16Arena should not be treated as offering an online money gaming service in respect of Feature A as described. A payout of a company-funded prize to a winning user is not, on these facts, a transaction towards payment for an online money gaming service.")
    para(doc, "5.3 Accordingly, Section 7 should not be engaged merely because the addressee bank or payment aggregator facilitates the payout of such company-funded prizes, provided the architecture and compliance controls described in this representation are actually implemented and maintained.")

    heading(doc, "6. Game Shop and Payments Caveat", 1)
    add_callout(
        doc,
        "Payments caveat",
        "This representation is not an RBI authorisation opinion. If the Game Shop or any stored-value arrangement permits use at third-party merchants, cash-out, transfer, or settlement beyond purchase of goods/services from 16Arena itself, a separate payments-law review is required before launch or continuation.",
        fill="FFF7ED",
        border=GOLD,
    )
    para(doc, "6.1 The Game Shop feature has not been relied upon as a stake for Feature A. It must remain technically and financially separated from tournaments.")
    para(doc, "6.2 I recommend that the addressees require from 16Arena a payments counsel note or equivalent confirmation regarding the Game Shop structure and any internal prize ledger before activating production payout flows.")

    heading(doc, "7. Client Compliance Commitments", 1)
    for item in [
        "maintain the no-fee, no-deposit, no-stake model for Feature A;",
        "route payouts only through lawful bank/payment partner rails after KYC and TDS checks;",
        "avoid manual Aadhaar collection/storage except through a lawful authorised KYC process;",
        "maintain TAN, TDS computation, quarterly filing and Form 16A workflows as advised by the CA;",
        "operate state restrictions, 18+ age gate, grievance redressal, user-safety and fair-play monitoring;",
        "notify the addressees in advance of any change that introduces paid access, user deposits, paid tokens, prize-to-shop transfer, wallet cash-out or other payment-linked change.",
    ]:
        bullet(doc, item)

    heading(doc, "8. Reliance and Limitations", 1)
    para(doc, "8.1 This representation is addressed only to the named bank and payment aggregator for diligence in relation to 16Arena. It is not an indemnity, guarantee, certification of technical implementation, tax opinion, GST opinion, PPI authorisation opinion, or audit report.")
    para(doc, "8.2 The addressees may form their own view and may require further documents, including product walkthroughs, PA agreements, bank account confirmations, CA/CS notes, code-level architecture certificates and OGAI correspondence.")
    para(doc, "8.3 Any departure from the facts assumed, any change in law, any OGAI direction or determination, or any contrary judicial pronouncement may alter this representation.")

    para(doc, "I remain available for any clarification required by the addressees.", before=6, after=16)
    para(doc, "Yours faithfully,", after=18)
    para(doc, "__________________________________", after=3)
    para(doc, "Shashwat", bold=True, color=NAVY, after=1)
    para(doc, "Advocate, Patna High Court and Supreme Court of India", after=10)

    heading(doc, "Enclosures", 2)
    for item in [
        "Final Internal Legal Determination Memorandum dated [insert] May 2026",
        "Application-cum-Representation proposed to be filed before OGAI",
        "Legal Opinion dated 20 May 2026",
        "Product architecture and payout flow materials",
        "CA and CS compliance notes, when finalised",
    ]:
        bullet(doc, item)

    path = OUT / "16Arena_Final_Counsels_Representation_Bank_PA.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    paths = [
        build_internal_memo(),
        build_ogai_application(),
        build_counsel_representation(),
    ]
    for p in paths:
        print(p)
